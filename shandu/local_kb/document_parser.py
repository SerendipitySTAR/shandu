import os
import logging
from typing import Callable, Dict, Any, Optional

try:
    import PyPDF2
except ImportError:
    logging.warning("PyPDF2 library not found. PDF parsing will not be available. "
                    "Please install it with 'pip install pypdf2'")
    PyPDF2 = None # type: ignore

try:
    import docx
except ImportError:
    logging.warning("python-docx library not found. DOCX parsing will not be available. "
                    "Please install it with 'pip install python-docx'")
    docx = None # type: ignore

# Configure logger for the module
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts text from a PDF file.

    Args:
        file_path: The path to the PDF file.

    Returns:
        The extracted text, or an empty string if extraction fails or PyPDF2 is not available.
    """
    if not PyPDF2:
        logger.error("PyPDF2 library is not installed. Cannot parse PDF file: %s", file_path)
        return ""
        
    text = []
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text.append(page.extract_text() or "") # Ensure None is handled
        logger.info("Successfully extracted text from PDF: %s", file_path)
        return "\n".join(text)
    except FileNotFoundError:
        logger.error("PDF file not found: %s", file_path)
        return ""
    except PyPDF2.errors.PdfReadError as e:
        logger.error("Error reading PDF file %s (possibly corrupted or password-protected): %s", file_path, e)
        return ""
    except Exception as e:
        logger.error("An unexpected error occurred while parsing PDF file %s: %s", file_path, e)
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file.

    Args:
        file_path: The path to the DOCX file.

    Returns:
        The extracted text, or an empty string if extraction fails or python-docx is not available.
    """
    if not docx:
        logger.error("python-docx library is not installed. Cannot parse DOCX file: %s", file_path)
        return ""

    text = []
    try:
        document = docx.Document(file_path)
        for para in document.paragraphs:
            text.append(para.text)
        logger.info("Successfully extracted text from DOCX: %s", file_path)
        return "\n".join(text)
    except FileNotFoundError:
        logger.error("DOCX file not found: %s", file_path)
        return ""
    except Exception as e: # docx library might raise various exceptions for corrupted files
        logger.error("An unexpected error occurred while parsing DOCX file %s: %s", file_path, e)
        return ""

def extract_text_from_txt(file_path: str) -> str:
    """
    Extracts text from a TXT file.

    Args:
        file_path: The path to the TXT file.

    Returns:
        The content of the file, or an empty string if an error occurs.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        logger.info("Successfully read text from TXT: %s", file_path)
        return content
    except FileNotFoundError:
        logger.error("TXT file not found: %s", file_path)
        return ""
    except UnicodeDecodeError as e:
        logger.error("Encoding error while reading TXT file %s (try a different encoding if not UTF-8): %s", file_path, e)
        return ""
    except Exception as e:
        logger.error("An unexpected error occurred while reading TXT file %s: %s", file_path, e)
        return ""

FILE_PARSERS: Dict[str, Callable[[str], str]] = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_txt,
}

def parse_document(file_path: str) -> Optional[str]:
    """
    Parses a document based on its file extension.

    Args:
        file_path: The path to the document.

    Returns:
        The extracted text as a string, or None if parsing fails or the file type is unsupported.
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file {file_path} does not exist.")
            
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()

        parser = FILE_PARSERS.get(file_extension)

        if parser:
            logger.info("Parsing document: %s using parser for %s", file_path, file_extension)
            return parser(file_path)
        else:
            logger.warning("Unsupported file type for document: %s (extension: %s)", file_path, file_extension)
            return None
    except FileNotFoundError as e:
        logger.error("File not found during parsing: %s", e)
        return None
    except Exception as e:
        logger.error("An unexpected error occurred in parse_document for file %s: %s", file_path, e)
        return None

if __name__ == '__main__':
    # Example Usage (for testing purposes)
    # Create dummy files for testing
    TEST_DIR = "test_docs"
    os.makedirs(TEST_DIR, exist_ok=True)

    DUMMY_TXT_FILE = os.path.join(TEST_DIR, "dummy.txt")
    DUMMY_DOCX_FILE = os.path.join(TEST_DIR, "dummy.docx")
    DUMMY_PDF_FILE = os.path.join(TEST_DIR, "dummy.pdf") # Requires a real PDF for actual text extraction
    UNSUPPORTED_FILE = os.path.join(TEST_DIR, "dummy.rtf")

    with open(DUMMY_TXT_FILE, "w", encoding="utf-8") as f:
        f.write("This is a test text file.\nIt has multiple lines.")

    if docx:
        try:
            doc = docx.Document()
            doc.add_paragraph("This is a test DOCX file.")
            doc.add_paragraph("It also has multiple lines of content.")
            doc.save(DUMMY_DOCX_FILE)
        except Exception as e:
            logger.warning("Could not create dummy DOCX file for testing: %s. python-docx might not be fully functional.", e)


    # Note: Creating a valid dummy PDF programmatically is complex with PyPDF2 (it's more for reading/manipulating).
    # For PDF testing, usually a small, actual PDF file would be used.
    # We'll simulate its presence for the parse_document logic.
    # To actually test PDF extraction, place a real PDF named dummy.pdf in the test_docs directory.
    if PyPDF2:
        # Create a very simple PDF if possible, or just acknowledge its absence for full test
        try:
            # PyPDF2 is not designed for creating PDFs from scratch easily.
            # This section is more illustrative. For a real test, provide a dummy.pdf.
            if not os.path.exists(DUMMY_PDF_FILE):
                 logger.info(f"To test PDF parsing, please place a valid PDF file at: {DUMMY_PDF_FILE}")

        except Exception as e:
            logger.warning("Could not create or check dummy PDF file for testing: %s", e)


    with open(UNSUPPORTED_FILE, "w", encoding="utf-8") as f:
        f.write("This is an unsupported RTF file.")

    print(f"--- Testing TXT ({DUMMY_TXT_FILE}) ---")
    txt_content = parse_document(DUMMY_TXT_FILE)
    if txt_content is not None:
        print(f"Extracted TXT Content:\n{txt_content}\n")
    
    if docx:
        print(f"--- Testing DOCX ({DUMMY_DOCX_FILE}) ---")
        docx_content = parse_document(DUMMY_DOCX_FILE)
        if docx_content is not None:
            print(f"Extracted DOCX Content:\n{docx_content}\n")
    else:
        print("Skipping DOCX test as python-docx is not available.\n")

    if PyPDF2 and os.path.exists(DUMMY_PDF_FILE):
        print(f"--- Testing PDF ({DUMMY_PDF_FILE}) ---")
        pdf_content = parse_document(DUMMY_PDF_FILE)
        if pdf_content is not None: # Will be empty string if PDF is not readable or dummy
            print(f"Extracted PDF Content:\n'{pdf_content}'\n") # Print quotes to show if it's empty
        elif os.path.exists(DUMMY_PDF_FILE):
             print(f"PDF content was None, but file exists. Check if it's a valid PDF and PyPDF2 can read it.\n")
    elif PyPDF2:
        print(f"Skipping PDF test as dummy PDF file '{DUMMY_PDF_FILE}' not found.\n")
    else:
        print("Skipping PDF test as PyPDF2 is not available.\n")

    print(f"--- Testing Unsupported File ({UNSUPPORTED_FILE}) ---")
    unsupported_content = parse_document(UNSUPPORTED_FILE)
    if unsupported_content is None:
        print("Correctly returned None for unsupported file type.\n")

    print(f"--- Testing Non-existent File ---")
    non_existent_content = parse_document("non_existent_file.txt")
    if non_existent_content is None:
        print("Correctly returned None for non-existent file.\n")

    # Clean up dummy files
    # import shutil
    # shutil.rmtree(TEST_DIR)
    # print(f"Cleaned up test directory: {TEST_DIR}")
