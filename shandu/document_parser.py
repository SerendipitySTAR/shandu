import os
from typing import Dict, Any, Optional
import chardet # For robust encoding detection in TXT files

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError as DocxPackageNotFoundError
except ImportError:
    Document = None 
    DocxPackageNotFoundError = None 

try:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError
except ImportError:
    PdfReader = None
    PdfReadError = None


def _get_filename_title(file_path: str) -> str:
    """Helper to get filename as title."""
    return os.path.basename(file_path)

def parse_txt(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Parses a TXT file.
    Returns a dictionary with text_content, title, and empty metadata.
    """
    if not os.path.exists(file_path):
        return None
    try:
        with open(file_path, 'rb') as f: # Read in binary mode for chardet
            raw_data = f.read()
        
        detected_encoding = chardet.detect(raw_data)['encoding']
        
        if detected_encoding:
            text_content = raw_data.decode(detected_encoding)
        else: # Fallback if chardet fails
            text_content = raw_data.decode('utf-8', errors='replace') # Try UTF-8 first

        return {
            "text_content": text_content,
            "title": _get_filename_title(file_path),
            "metadata": {}
        }
    except FileNotFoundError:
        return None
    except UnicodeDecodeError:
        try:
            # Fallback to latin-1 if initial decoding fails
            with open(file_path, 'r', encoding='latin-1') as f:
                text_content = f.read()
            return {
                "text_content": text_content,
                "title": _get_filename_title(file_path),
                "metadata": {}
            }
        except Exception: # Catch any other error during fallback
            return None
    except Exception: # Catch-all for other read errors
        return None

def parse_docx(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Parses a DOCX file using python-docx.
    Returns a dictionary with text_content, title, and metadata.
    """
    if Document is None:
        # python-docx is not installed
        return {"text_content": "", "title": _get_filename_title(file_path), "metadata": {"error": "python-docx library not installed."}}

    if not os.path.exists(file_path):
        return None
    try:
        doc = Document(file_path)
        text_content = "\n".join([para.text for para in doc.paragraphs])
        
        metadata = {}
        core_props = doc.core_properties
        if core_props:
            metadata["author"] = core_props.author or None
            # Ensure datetime objects are converted to string (ISO format)
            if core_props.created:
                metadata["created"] = core_props.created.isoformat() if hasattr(core_props.created, 'isoformat') else str(core_props.created)
            else:
                metadata["created"] = None
            if core_props.last_modified_by: # DOCX stores last_modified_by not last_modified
                 metadata["last_modified_by"] = core_props.last_modified_by or None
            if core_props.modified:
                 metadata["modified"] = core_props.modified.isoformat() if hasattr(core_props.modified, 'isoformat') else str(core_props.modified)
            else:
                metadata["modified"] = None


        doc_title = core_props.title if core_props and core_props.title else None
        
        return {
            "text_content": text_content,
            "title": doc_title or _get_filename_title(file_path),
            "metadata": metadata
        }
    except DocxPackageNotFoundError: # Handles issues like file not being a valid DOCX zip
        return None
    except Exception: # Catch-all for other parsing errors
        return None

def parse_pdf(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Parses a PDF file using pypdf.
    Returns a dictionary with text_content, title, and metadata.
    """
    if PdfReader is None:
        # pypdf is not installed
        return {"text_content": "", "title": _get_filename_title(file_path), "metadata": {"error": "pypdf library not installed."}}

    if not os.path.exists(file_path):
        return None
    try:
        reader = PdfReader(file_path)
        text_content = ""
        for page in reader.pages:
            text_content += page.extract_text() or "" # Add space or newline if desired
        
        metadata = {}
        doc_info = reader.metadata
        if doc_info:
            metadata["author"] = doc_info.author or None
            metadata["title"] = doc_info.title or None # PDF metadata title
            # PDF dates are often in a specific string format, D:YYYYMMDDHHMMSSZ...'
            # For simplicity, we'll store them as strings if they exist.
            # More sophisticated parsing can be added if specific datetime objects are needed.
            if doc_info.creation_date:
                 metadata["creation_date"] = str(doc_info.creation_date)
            else:
                metadata["creation_date"] = None
            if doc_info.modification_date:
                metadata["mod_date"] = str(doc_info.modification_date)
            else:
                metadata["mod_date"] = None
        
        pdf_meta_title = metadata.get("title")

        return {
            "text_content": text_content,
            "title": pdf_meta_title or _get_filename_title(file_path),
            "metadata": metadata
        }
    except PdfReadError: # Specific error for pypdf
        return None
    except Exception: # Catch-all for other parsing errors
        return None

def parse_document(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Facade function to parse a document based on its extension.
    """
    if not isinstance(file_path, str): # Basic type check
        return None
        
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".txt":
        return parse_txt(file_path)
    elif ext == ".docx":
        if Document is None: # Check if library was imported
             return {"text_content": "", "title": _get_filename_title(file_path), "metadata": {"error": "python-docx library not installed, cannot parse .docx"}}
        return parse_docx(file_path)
    elif ext == ".pdf":
        if PdfReader is None: # Check if library was imported
            return {"text_content": "", "title": _get_filename_title(file_path), "metadata": {"error": "pypdf library not installed, cannot parse .pdf"}}
        return parse_pdf(file_path)
    else:
        # Optionally, handle unsupported file types
        return {"text_content": "", "title": _get_filename_title(file_path), "metadata": {"error": f"Unsupported file type: {ext}"}}

# Example Usage (can be removed or kept for testing)
if __name__ == '__main__':
    # Create dummy files for testing
    with open("dummy.txt", "w", encoding="utf-8") as f:
        f.write("This is a test TXT file. It has some text content.")
    
    if Document is not None:
        doc = Document()
        doc.add_paragraph("This is a test DOCX file.")
        doc.core_properties.title = "Dummy DOCX Title"
        doc.core_properties.author = "Test Author"
        doc.save("dummy.docx")

    # Note: Creating a dummy PDF programmatically is more complex,
    # so we'll assume one exists for testing or test manually.
    # For CI/automated testing, a small, simple PDF should be included.

    print("--- Testing TXT Parser ---")
    txt_data = parse_document("dummy.txt")
    if txt_data:
        print(f"Title: {txt_data['title']}")
        print(f"Content Preview: {txt_data['text_content'][:50]}...")
        print(f"Metadata: {txt_data['metadata']}")
    else:
        print("Failed to parse TXT.")

    print("\n--- Testing DOCX Parser ---")
    if Document is not None:
        docx_data = parse_document("dummy.docx")
        if docx_data:
            print(f"Title: {docx_data['title']}")
            print(f"Content Preview: {docx_data['text_content'][:50]}...")
            print(f"Metadata: {docx_data['metadata']}")
        else:
            print("Failed to parse DOCX (or file doesn't exist).")
    else:
        print("python-docx library not installed, skipping DOCX test.")

    print("\n--- Testing PDF Parser (Manual) ---")
    # Create a dummy PDF or point to an existing small PDF for testing
    # e.g., pdf_data = parse_document("dummy.pdf")
    # if pdf_data:
    #     print(f"Title: {pdf_data['title']}")
    #     print(f"Content Preview: {pdf_data['text_content'][:50]}...")
    #     print(f"Metadata: {pdf_data['metadata']}")
    # else:
    #     print("Failed to parse PDF (or file doesn't exist/pypdf not installed).")
    
    print("\n--- Testing Unsupported Type ---")
    unsupported_data = parse_document("dummy.xyz")
    if unsupported_data and "error" in unsupported_data["metadata"]:
        print(f"Error for .xyz: {unsupported_data['metadata']['error']}")
    else:
        print("Unexpected result for unsupported type.")

    # Clean up dummy files
    if os.path.exists("dummy.txt"):
        os.remove("dummy.txt")
    if os.path.exists("dummy.docx"):
        os.remove("dummy.docx")
